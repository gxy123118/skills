from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"


def build_flowchart_pict() -> str:
    return f"""
<w:pict {nsdecls('w')} xmlns:v="{V_NS}" xmlns:o="{O_NS}">
  <v:group coordorigin="0,0" coordsize="720,180"
    style="position:relative;width:468pt;height:132pt;">
    <v:rect id="startBox" fillcolor="#EAF2F8" strokecolor="#4F81BD"
      style="position:absolute;left:24;top:54;width:150;height:54;">
      <v:textbox inset="6pt,4pt,6pt,4pt">
        <w:txbxContent>
          <w:p>
            <w:pPr><w:jc w:val="center"/></w:pPr>
            <w:r><w:rPr><w:sz w:val="20"/><w:b/></w:rPr><w:t>Start</w:t></w:r>
          </w:p>
        </w:txbxContent>
      </v:textbox>
    </v:rect>

    <v:line id="arrowOne" from="174,81" to="282,81"
      strokecolor="#4F81BD" style="position:absolute;">
      <v:stroke weight="2pt" endarrow="block"/>
    </v:line>

    <v:roundrect id="processBox" arcsize="10%" fillcolor="#FDFEFE" strokecolor="#5D6D7E"
      style="position:absolute;left:282;top:42;width:156;height:78;">
      <v:textbox inset="6pt,4pt,6pt,4pt">
        <w:txbxContent>
          <w:p>
            <w:pPr><w:jc w:val="center"/></w:pPr>
            <w:r><w:rPr><w:sz w:val="20"/><w:b/></w:rPr><w:t>Review</w:t></w:r>
          </w:p>
          <w:p>
            <w:pPr><w:jc w:val="center"/></w:pPr>
            <w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>Check content</w:t></w:r>
          </w:p>
        </w:txbxContent>
      </v:textbox>
    </v:roundrect>

    <v:line id="arrowTwo" from="438,81" to="546,81"
      strokecolor="#4F81BD" style="position:absolute;">
      <v:stroke weight="2pt" endarrow="block"/>
    </v:line>

    <v:rect id="endBox" fillcolor="#E8F8F5" strokecolor="#2E8B57"
      style="position:absolute;left:546;top:54;width:150;height:54;">
      <v:textbox inset="6pt,4pt,6pt,4pt">
        <w:txbxContent>
          <w:p>
            <w:pPr><w:jc w:val="center"/></w:pPr>
            <w:r><w:rPr><w:sz w:val="20"/><w:b/></w:rPr><w:t>Done</w:t></w:r>
          </w:p>
        </w:txbxContent>
      </v:textbox>
    </v:rect>
  </v:group>
</w:pict>
""".strip()


def main():
    output_path = r"C:\Users\Administrator\Desktop\docx-native-shapes-example.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCX Native Shapes Example")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("This flowchart is built with Word native VML shapes, not an embedded image.")
    subrun.font.size = Pt(10.5)

    doc.add_paragraph()

    flow_para = doc.add_paragraph()
    flow_run = flow_para.add_run()
    flow_run._r.append(parse_xml(build_flowchart_pict()))

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.style = doc.styles["Normal"]
    note.add_run("Example flow: ").bold = True
    note.add_run("Start -> Review -> Done")

    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
